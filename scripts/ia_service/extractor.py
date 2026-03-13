"""
Extractor de texto desde archivos — usado por el RAG.

Soporta: PDF, DOCX, DOC, XLSX, XLS, TXT, CSV, MD, imágenes (via Gemini Vision).
"""
import os
import tempfile


def extraer_texto(ruta_archivo: str, nombre_archivo: str = '') -> str:
    """
    Extrae texto de un archivo dado su path temporal.
    Detecta el tipo por extensión y contenido.
    Retorna string de texto limpio.
    """
    ext = os.path.splitext(nombre_archivo or ruta_archivo)[1].lower().strip('.')
    tamanio = os.path.getsize(ruta_archivo)

    if tamanio == 0:
        raise ValueError("Archivo vacío")
    if tamanio > 50 * 1024 * 1024:  # 50 MB
        raise ValueError("Archivo demasiado grande (máx. 50 MB)")

    if ext == 'pdf':
        return _extraer_pdf(ruta_archivo)
    elif ext in ('docx',):
        return _extraer_docx(ruta_archivo)
    elif ext in ('doc',):
        return _extraer_doc(ruta_archivo)
    elif ext in ('xlsx', 'xls', 'xlsm'):
        return _extraer_excel(ruta_archivo)
    elif ext in ('csv',):
        return _extraer_csv(ruta_archivo)
    elif ext in ('txt', 'md', 'rst', 'text'):
        return _extraer_texto_plano(ruta_archivo)
    elif ext in ('jpg', 'jpeg', 'png', 'webp', 'gif', 'bmp'):
        return _extraer_imagen(ruta_archivo, nombre_archivo)
    else:
        # Intentar como texto plano (para .json, .xml, .html, etc.)
        try:
            return _extraer_texto_plano(ruta_archivo)
        except Exception:
            raise ValueError(f"Formato no soportado: .{ext}")


# ─────────────────────────────────────────────────────────
# Extractores por tipo
# ─────────────────────────────────────────────────────────

def _extraer_pdf(ruta: str) -> str:
    """Extrae texto de PDF usando pymupdf (fitz) — más robusto que pypdf."""
    try:
        import fitz  # pymupdf
        doc = fitz.open(ruta)
        paginas = []
        for i, pag in enumerate(doc):
            texto = pag.get_text().strip()
            if texto:
                paginas.append(f"[Página {i+1}]\n{texto}")
        doc.close()
        if not paginas:
            raise ValueError("PDF sin texto extraíble (puede ser imagen escaneada)")
        return '\n\n'.join(paginas)
    except ImportError:
        # Fallback a pypdf
        import pypdf
        reader = pypdf.PdfReader(ruta)
        paginas = []
        for i, pag in enumerate(reader.pages):
            texto = pag.extract_text() or ''
            if texto.strip():
                paginas.append(f"[Página {i+1}]\n{texto.strip()}")
        if not paginas:
            raise ValueError("PDF sin texto extraíble")
        return '\n\n'.join(paginas)


def _extraer_docx(ruta: str) -> str:
    """Extrae texto de archivos DOCX (Word moderno)."""
    import docx
    doc = docx.Document(ruta)
    partes = []
    for para in doc.paragraphs:
        texto = para.text.strip()
        if texto:
            partes.append(texto)
    # También tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celdas:
                partes.append(' | '.join(celdas))
    if not partes:
        raise ValueError("Documento Word vacío o sin texto")
    return '\n\n'.join(partes)


def _extraer_doc(ruta: str) -> str:
    """Extrae texto de archivos DOC (Word antiguo) usando antiword si está disponible."""
    import subprocess
    try:
        resultado = subprocess.run(
            ['antiword', ruta], capture_output=True, text=True, timeout=30
        )
        if resultado.returncode == 0 and resultado.stdout.strip():
            return resultado.stdout.strip()
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    # Fallback: intentar leer como DOCX (a veces funciona)
    try:
        return _extraer_docx(ruta)
    except Exception:
        raise ValueError("No se pudo extraer texto del .doc — usa .docx para mejor soporte")


def _extraer_excel(ruta: str) -> str:
    """Extrae datos de Excel como texto tabular."""
    import openpyxl
    wb = openpyxl.load_workbook(ruta, read_only=True, data_only=True)
    partes = []
    for nombre_hoja in wb.sheetnames:
        hoja = wb[nombre_hoja]
        filas_hoja = []
        for fila in hoja.iter_rows(values_only=True):
            celdas = [str(c).strip() if c is not None else '' for c in fila]
            if any(c for c in celdas):  # fila no vacía
                filas_hoja.append(' | '.join(celdas))
        if filas_hoja:
            partes.append(f"## Hoja: {nombre_hoja}\n" + '\n'.join(filas_hoja))
    wb.close()
    if not partes:
        raise ValueError("Excel sin datos")
    return '\n\n'.join(partes)


def _extraer_csv(ruta: str) -> str:
    """Extrae CSV como texto."""
    import csv
    filas = []
    encodings = ['utf-8', 'latin-1', 'cp1252']
    for enc in encodings:
        try:
            with open(ruta, encoding=enc, newline='') as f:
                reader = csv.reader(f)
                for fila in reader:
                    if any(c.strip() for c in fila):
                        filas.append(' | '.join(c.strip() for c in fila))
            break
        except UnicodeDecodeError:
            continue
    if not filas:
        raise ValueError("CSV vacío o sin datos")
    return '\n'.join(filas)


def _extraer_texto_plano(ruta: str) -> str:
    """Lee archivos de texto con detección automática de encoding."""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
    for enc in encodings:
        try:
            with open(ruta, encoding=enc) as f:
                texto = f.read().strip()
            if texto:
                return texto
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError("No se pudo decodificar el archivo de texto")


def _extraer_imagen(ruta: str, nombre: str) -> str:
    """
    Extrae texto de imagen usando Gemini Vision.
    Describe la imagen y extrae cualquier texto visible.
    """
    from .config import get_local_conn
    import base64
    import urllib.request
    import json as json_mod

    # Obtener API key de Gemini
    conn = get_local_conn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT api_key FROM ia_agentes WHERE slug = 'gemini-pro' AND activo = 1 LIMIT 1"
            )
            row = cur.fetchone()
    finally:
        conn.close()

    if not row or not row.get('api_key'):
        raise ValueError("No hay API key de Gemini disponible para procesar imágenes")

    api_key = row['api_key']

    # Leer imagen en base64
    with open(ruta, 'rb') as f:
        datos = f.read()
    b64 = base64.b64encode(datos).decode()

    ext = os.path.splitext(nombre)[1].lower()
    mime_map = {'.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                '.png': 'image/png', '.webp': 'image/webp',
                '.gif': 'image/gif', '.bmp': 'image/png'}
    mime = mime_map.get(ext, 'image/jpeg')

    payload = json_mod.dumps({
        "contents": [{
            "parts": [
                {"text": (
                    "Analiza esta imagen y extrae TODA la información que contiene. "
                    "Si hay texto, transcríbelo completo. "
                    "Si es una tabla o gráfico, describe los datos. "
                    "Si es una imagen sin texto, describe su contenido en detalle. "
                    "Responde en el mismo idioma del contenido."
                )},
                {"inline_data": {"mime_type": mime, "data": b64}}
            ]
        }]
    }).encode()

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}"
    req = urllib.request.Request(url, data=payload, headers={'Content-Type': 'application/json'})
    with urllib.request.urlopen(req, timeout=60) as resp:
        resultado = json_mod.loads(resp.read())

    texto = resultado['candidates'][0]['content']['parts'][0]['text']
    return f"[Imagen: {nombre}]\n{texto}"
